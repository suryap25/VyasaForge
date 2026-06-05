"""DOCX compiler for Milestone M9."""

from __future__ import annotations

from pathlib import Path

from src.handbook import resolve_chapter

OUTPUT_PATH = Path("output/AppSec_Authentication_Authorization_Handbook_Phase1.docx")
TITLE = "AppSec Authentication & Authorization Handbook v2.0"
SUBTITLE = "Phase 1: Foundations & JWT"


def add_markdown_line(document: object, line: str) -> None:
    """Add a supported Markdown line to a DOCX document."""
    if line.startswith("### "):
        document.add_heading(line[4:].strip(), level=3)
    elif line.startswith("## "):
        document.add_heading(line[3:].strip(), level=2)
    elif line.startswith("# "):
        document.add_heading(line[2:].strip(), level=1)
    elif line.startswith("- "):
        document.add_paragraph(line[2:].strip(), style="List Bullet")
    elif line.strip():
        document.add_paragraph(line.strip())


def compile_docx(chapters: list[int]) -> Path:
    """Compile final Markdown chapters into a DOCX file."""
    chapter_paths = [resolve_chapter(chapter).final_path for chapter in chapters]
    for chapter_path in chapter_paths:
        if not chapter_path.exists():
            raise FileNotFoundError(f"Missing final chapter: {chapter_path}")

    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required. Install project dependencies with `pip install -e .`.") from exc

    document = Document()
    document.add_heading(TITLE, level=0)
    document.add_paragraph(SUBTITLE)
    document.add_page_break()

    for index, chapter_path in enumerate(chapter_paths):
        if index > 0:
            document.add_page_break()

        markdown = chapter_path.read_text(encoding="utf-8")
        for line in markdown.splitlines():
            add_markdown_line(document, line)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH
